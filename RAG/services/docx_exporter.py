"""Convert edited workspace Markdown to a formatted Word document (.docx).

Translates headings, lists, basic inline styling (bold, italic, underline, strike,
highlight), inline code, external links, tables, and embeds local figures correctly.

Every export is built on top of `templates/socradar_template.docx` so the SOCRadar
cover watermark, header/footer and page geometry carry over automatically; only
the body content is replaced. Heading and body runs are styled directly (rather
than via the template's generic Heading styles) to match the brand look used in
the template itself: Red Hat Display ExtraBold headings in coral, Inter body copy
in dark gray.
"""
from pathlib import Path
import markdown as md_lib
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

from RAG.services import paths

_TEMPLATE_DOCX = Path(__file__).resolve().parent / "templates" / "socradar_template.docx"

# Brand colours sampled from SOCRadar Template.docx (cover watermark + heading runs).
_BRAND_ACCENT = RGBColor(0xFF, 0x45, 0x62)   # coral — heading colour in the template
_BRAND_NAVY = RGBColor(0x19, 0x19, 0x38)     # deep navy from the cover banner
_BRAND_TEXT = RGBColor(0x43, 0x43, 0x43)     # body copy gray used in the template
_BRAND_MUTED = RGBColor(0x66, 0x66, 0x66)

_HEADING_FONT = "Red Hat Display ExtraBold"
_BODY_FONT = "Inter"

# Point sizes per heading level, descending — matches the template's own
# heading/subheading proportions (its "Header" ~19pt down to "Header 2" ~13pt).
_HEADING_SIZES = {1: 22, 2: 17, 3: 14, 4: 12, 5: 11, 6: 11}


def _new_document() -> Document:
    """Start from the SOCRadar template (keeps its watermark header/footer/page
    geometry) with the placeholder body content stripped out, falling back to a
    blank document if the template is missing."""
    if _TEMPLATE_DOCX.exists():
        doc = Document(str(_TEMPLATE_DOCX))
        body = doc.element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
        return doc

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _BODY_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = _BRAND_TEXT
    return doc


def _style_heading_run(run, level: int):
    run.font.name = _HEADING_FONT
    run.font.size = Pt(_HEADING_SIZES.get(level, 11))
    run.font.color.rgb = _BRAND_ACCENT
    run.font.bold = False  # the ExtraBold weight is baked into the font itself


def _style_body_run(run):
    if not run.font.name:
        run.font.name = _BODY_FONT
    if run.font.color.rgb is None:
        run.font.color.rgb = _BRAND_TEXT


def _shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tcPr.append(shd)


def _set_table_borders(table, hex_color: str = "D9B3B9"):
    """Draw plain grid borders directly (the SOCRadar template has no built-in
    'Table Grid' style since it ships with zero tables)."""
    tblPr = table._tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tblPr.makeelement(
            qn(f"w:{edge}"),
            {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "0", qn("w:color"): hex_color},
        )
        borders.append(el)
    tblPr.append(borders)


def render(source: str) -> Path:
    """Render workspace markdown to a DOCX file and return its path."""
    md_path = paths.workspace_md_path(source)
    if not md_path.exists():
        raise FileNotFoundError(f"Workspace markdown not found: {md_path}")

    paths.ensure_dirs()
    text = md_path.read_text(encoding="utf-8")

    # Render markdown to HTML body
    html_body = md_lib.markdown(
        text,
        extensions=["tables", "fenced_code", "toc", "sane_lists", "nl2br"],
    )

    doc = _new_document()

    soup = BeautifulSoup(html_body, "html.parser")

    # Walk children of the HTML body
    for el in soup.children:
        if el.name is None:
            continue

        # Headings
        if el.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(el.name[1])
            p = doc.add_paragraph()
            p.style = doc.styles[f"Heading {level}"]
            p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(el.text)
            _style_heading_run(run, level)
            if level == 1:
                # Coral rule under H1s, echoing the template's section dividers.
                p.paragraph_format.space_after = Pt(4)
                pPr = p._p.get_or_add_pPr()
                pBdr = pPr.makeelement(qn("w:pBdr"), {})
                bottom = pPr.makeelement(
                    qn("w:bottom"),
                    {qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "4", qn("w:color"): "FF4562"},
                )
                pBdr.append(bottom)
                pPr.append(pBdr)

        # Paragraphs
        elif el.name == "p":
            p = doc.add_paragraph()
            _process_inline(p, el, source)
            for run in p.runs:
                _style_body_run(run)

        # Lists (ul, ol) — the SOCRadar template ships with no List Bullet/Number
        # styles (it has zero lists), so indent + a manual marker instead of
        # relying on a named style that would raise KeyError.
        elif el.name in ("ul", "ol"):
            for idx, li in enumerate(el.find_all("li", recursive=False), start=1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                marker = "•  " if el.name == "ul" else f"{idx}.  "
                marker_run = p.add_run(marker)
                _style_body_run(marker_run)
                _process_inline(p, li, source)
                for run in p.runs:
                    _style_body_run(run)

        # Blockquotes
        elif el.name == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _process_inline(p, el, source)
            # Make the blockquote italic
            for run in p.runs:
                _style_body_run(run)
                run.italic = True

        # Tables
        elif el.name == "table":
            rows = el.find_all("tr")
            if not rows:
                continue

            # Compute max columns
            max_cols = 1
            for row in rows:
                cols = len(row.find_all(["td", "th"]))
                if cols > max_cols:
                    max_cols = cols

            table = doc.add_table(rows=len(rows), cols=max_cols)
            _set_table_borders(table)

            for r_idx, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                is_header_row = r_idx == 0
                for c_idx, cell in enumerate(cells):
                    if c_idx >= max_cols:
                        continue
                    doc_cell = table.cell(r_idx, c_idx)
                    doc_cell.text = cell.text
                    for p in doc_cell.paragraphs:
                        for run in p.runs:
                            run.font.name = _BODY_FONT
                            run.font.size = Pt(10)
                            if is_header_row:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            else:
                                run.font.color.rgb = _BRAND_TEXT
                    if is_header_row:
                        _shade_cell(doc_cell, "FF4562")
                    elif r_idx % 2 == 0:
                        _shade_cell(doc_cell, "FFF4F6")  # pale coral zebra stripe

        # Fenced code/pre
        elif el.name == "pre":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            code = el.find("code")
            txt = code.text if code else el.text
            run = p.add_run(txt)
            run.font.name = "Courier New"
            run.font.size = Pt(10)

        # Horizontal rules
        elif el.name == "hr":
            doc.add_paragraph("---")

        # Div elements (cover-page, page-break, chart, KPI card, TOC)
        elif el.name == "div":
            classes = el.get("class", [])
            if "page-break" in classes:
                doc.add_page_break()
            elif "cover-page" in classes:
                title_el = el.find(class_="cover-title")
                sub_el = el.find(class_="cover-subtitle")
                meta_el = el.find(class_="cover-meta")

                title_text = title_el.text.strip() if title_el else "Untitled Report"
                sub_text = sub_el.text.strip() if sub_el else ""
                meta_text = meta_el.text.strip() if meta_el else ""

                for _ in range(4):
                    doc.add_paragraph()

                p_title = doc.add_paragraph()
                p_title.alignment = 1  # Center
                run_title = p_title.add_run(title_text)
                run_title.font.name = _HEADING_FONT
                run_title.font.size = Pt(28)
                run_title.font.color.rgb = _BRAND_ACCENT

                if sub_text:
                    p_sub = doc.add_paragraph()
                    p_sub.alignment = 1
                    run_sub = p_sub.add_run(sub_text)
                    run_sub.font.name = _BODY_FONT
                    run_sub.font.size = Pt(16)
                    run_sub.font.italic = True
                    run_sub.font.color.rgb = _BRAND_MUTED

                for _ in range(6):
                    doc.add_paragraph()

                if meta_text:
                    p_meta = doc.add_paragraph()
                    p_meta.alignment = 1
                    run_meta = p_meta.add_run(meta_text)
                    run_meta.font.name = _BODY_FONT
                    run_meta.font.size = Pt(11)
                    run_meta.font.color.rgb = _BRAND_TEXT

                doc.add_page_break()
            elif "report-chart-container" in classes:
                img = el.find("img")
                if img:
                    p = doc.add_paragraph()
                    _process_inline(p, el, source)
            elif "kpi-card" in classes:
                val_el = el.find(class_="kpi-value")
                lbl_el = el.find(class_="kpi-label")
                val_text = val_el.text.strip() if val_el else ""
                lbl_text = lbl_el.text.strip() if lbl_el else ""
                if val_text or lbl_text:
                    p = doc.add_paragraph()
                    run = p.add_run(f"★ {val_text} — {lbl_text}")
                    run.font.name = _BODY_FONT
                    run.font.bold = True
                    run.font.color.rgb = _BRAND_ACCENT
            else:
                p = doc.add_paragraph()
                _process_inline(p, el, source)
                for run in p.runs:
                    _style_body_run(run)

    out_path = paths.workspace_docx_path(source)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _process_inline(p, el, source: str):
    """Recursively process HTML children within an element and append runs to paragraph p."""
    for child in el.contents:
        if isinstance(child, str):
            # Text node
            if child.strip() == "" and not child.startswith(" ") and not child.endswith(" "):
                continue
            p.add_run(child)
        elif child.name in ("strong", "b"):
            run = p.add_run(child.text)
            run.bold = True
        elif child.name in ("em", "i"):
            run = p.add_run(child.text)
            run.italic = True
        elif child.name == "u":
            run = p.add_run(child.text)
            run.underline = True
        elif child.name == "mark":
            run = p.add_run(child.text)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif child.name in ("strike", "del", "s"):
            run = p.add_run(child.text)
            run.font.strike = True
        elif child.name == "code":
            run = p.add_run(child.text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif child.name == "a":
            # Render a blue/underlined text followed by the url
            run = p.add_run(child.text)
            run.font.underline = True
            run.font.color.rgb = RGBColor(0x19, 0x19, 0x38)
            href = child.get("href", "")
            if href:
                p.add_run(f" ({href})")
        elif child.name == "img":
            src = child.get("src", "")
            if src.startswith("/images/workspace/") or src.startswith("/images/originals/"):
                parts = src.strip("/").split("/")
                if len(parts) >= 4:
                    channel = parts[1]
                    stem = parts[2]
                    img_name = parts[3]
                    img_path = paths.image_path(channel, stem, img_name)
                    if img_path and img_path.exists():
                        # Add image on a new run
                        try:
                            p.add_run().add_picture(str(img_path), width=Inches(4.5))
                        except Exception as e:
                            print(f"[DOCXExporter] Failed to embed image {img_path}: {e}")
        else:
            # Fallback recurse
            _process_inline(p, child, source)
